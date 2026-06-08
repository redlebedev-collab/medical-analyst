"""
render.py — рендер problem list из YAML в Markdown и DOCX.

Принцип «единый источник»:
    problem_list.yaml ──render──▶ problem_list.md
                                  + problem_list.docx (альбомная ориентация)

Структура YAML:
    meta:
      patient_id: ...
      generated_at: ...
      disclaimer: ...
    problems:
      - id: P001
        priority: P1
        title: Пограничное АД
        data: "Среднее 138/88 (7-дн), тренд вверх с 2021"
        protocol: "ESC/ESH 2023, Class I Level A: целевое <130/80 при низком риске"
        action: "Суточное мониторирование, консультация кардиолога"
        specialist: Кардиолог
        analyses: [ОАМ, креатинин, калий]
        investigations: [СМАД, ЭКГ]
    not_doing:
      - item: Скрининг рака лёгкого
        reason: "Нет анамнеза курения; Grade D USPSTF"
    open_questions:
      - priority: P1
        question: "Уточнить семейный анамнез по АД"
        if_then: "Если у родственников 1-й линии — скорректировать риск"

Использование:
    python render.py                              # рендер из problem_list.yaml
    python render.py path/to/my_plan.yaml         # указать файл явно
"""

import sys
from datetime import datetime
from pathlib import Path

import yaml


# ---------------------------------------------------------------------------
# Загрузка YAML
# ---------------------------------------------------------------------------

def load_plan(yaml_path: Path) -> dict:
    with yaml_path.open(encoding="utf-8") as f:
        return yaml.safe_load(f)


# ---------------------------------------------------------------------------
# Рендер в Markdown
# ---------------------------------------------------------------------------

def render_markdown(plan: dict) -> str:
    meta = plan.get("meta", {})
    problems = plan.get("problems", [])
    not_doing = plan.get("not_doing", [])
    open_questions = plan.get("open_questions", [])

    lines = []

    # Заголовок
    lines.append(f"# Список проблем пациента")
    lines.append(f"")
    if meta.get("generated_at"):
        lines.append(f"_Дата: {meta['generated_at'][:10]}_")
    lines.append("")

    # Дисклеймер
    disclaimer = meta.get(
        "disclaimer",
        "**Документ подготовлен аналитическим инструментом, не врачом. "
        "Все выводы требуют обсуждения с лечащим специалистом.**"
    )
    lines.append(f"> {disclaimer}")
    lines.append("")

    # Таблица проблем
    lines.append("## Список проблем")
    lines.append("")
    lines.append("| № | Приоритет | Проблема | Данные / тренды | Протокол (гайдлайн) | Действие | Специалист |")
    lines.append("|---|-----------|----------|-----------------|---------------------|----------|------------|")

    all_analyses = set()
    all_investigations = set()
    all_specialists = set()

    for p in problems:
        priority = p.get("priority", "P2")
        priority_fmt = {"P1": "🔴 P1", "P2": "🟡 P2", "P3": "🟢 P3"}.get(priority, priority)
        lines.append(
            f"| {p.get('id', '')} "
            f"| {priority_fmt} "
            f"| {p.get('title', '')} "
            f"| {p.get('data', '')} "
            f"| {p.get('protocol', '')} "
            f"| {p.get('action', '')} "
            f"| {p.get('specialist', '')} |"
        )
        # Собираем для сводки
        for a in p.get("analyses", []):
            all_analyses.add(a)
        for i in p.get("investigations", []):
            all_investigations.add(i)
        sp = p.get("specialist", "")
        if sp:
            all_specialists.add(sp)

    lines.append("")

    # Сводка
    lines.append("## Сводка на визит")
    lines.append("")

    if all_analyses:
        lines.append("### Анализы (один забор крови)")
        for a in sorted(all_analyses):
            lines.append(f"- {a}")
        lines.append("")

    if all_investigations:
        lines.append("### Исследования")
        for i in sorted(all_investigations):
            lines.append(f"- {i}")
        lines.append("")

    if all_specialists:
        lines.append("### Специалисты")
        for s in sorted(all_specialists):
            lines.append(f"- {s}")
        lines.append("")

    # Что не делаем
    if not_doing:
        lines.append("## Что не делаем и почему")
        lines.append("")
        for nd in not_doing:
            item = nd.get("item", "")
            reason = nd.get("reason", "")
            lines.append(f"- **{item}** — {reason}")
        lines.append("")

    # Открытые вопросы / развилки
    if open_questions:
        lines.append("## Открытые вопросы и развилки")
        lines.append("")
        for oq in open_questions:
            priority = oq.get("priority", "P2")
            question = oq.get("question", "")
            if_then = oq.get("if_then", "")
            lines.append(f"**[{priority}]** {question}")
            if if_then:
                lines.append(f"  → {if_then}")
        lines.append("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Рендер в DOCX (альбомная ориентация для широкой таблицы)
# ---------------------------------------------------------------------------

def render_docx(markdown_text: str, output_path: Path) -> None:
    """
    Конвертирует Markdown → DOCX через python-docx.
    Устанавливает альбомную ориентацию страницы.
    """
    try:
        from docx import Document
        from docx.shared import Cm, Pt
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("⚠ python-docx не установлен: pip install python-docx")
        return

    doc = Document()

    # Альбомная ориентация (A4 landscape)
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Стиль по умолчанию
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(9)

    # Простой парсинг Markdown → параграфы и таблицы в DOCX
    lines = markdown_text.split("\n")
    i = 0
    table_header = None
    table_rows = []

    def flush_table():
        nonlocal table_header, table_rows
        if table_header is None:
            return
        cols = len(table_header)
        if not table_rows:
            table_header = None
            return

        tbl = doc.add_table(rows=1 + len(table_rows), cols=cols)
        tbl.style = "Table Grid"

        # Заголовок
        for j, cell_text in enumerate(table_header):
            cell = tbl.rows[0].cells[j]
            cell.text = cell_text.strip()
            run = cell.paragraphs[0].runs
            if run:
                run[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(8) if run else None

        # Данные
        for ri, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                if j < cols:
                    tbl.rows[ri + 1].cells[j].text = cell_text.strip()

        table_header = None
        table_rows = []

    while i < len(lines):
        line = lines[i]

        # Заголовки
        if line.startswith("# "):
            flush_table()
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            flush_table()
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            flush_table()
            p = doc.add_heading(line[4:], level=3)

        # Таблица
        elif line.startswith("|"):
            parts = [p.strip() for p in line.strip("|").split("|")]
            # Строка разделителей
            if all(set(p.replace("-", "").replace(":", "")) <= {""} for p in parts):
                i += 1
                continue
            if table_header is None:
                table_header = parts
            else:
                table_rows.append(parts)

        # Цитата (дисклеймер)
        elif line.startswith("> "):
            flush_table()
            p = doc.add_paragraph(line[2:])
            p.style = doc.styles["Normal"]
            run = p.runs[0] if p.runs else p.add_run("")
            run.italic = True

        # Маркированный список
        elif line.startswith("- "):
            flush_table()
            doc.add_paragraph(line[2:], style="List Bullet")

        # Обычный текст
        elif line.strip():
            flush_table()
            doc.add_paragraph(line)

        # Пустая строка
        else:
            flush_table()

        i += 1

    flush_table()

    doc.save(output_path)
    print(f"✓ DOCX сохранён: {output_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) > 1:
        yaml_path = Path(sys.argv[1])
    else:
        yaml_path = Path(__file__).parent / "problem_list.yaml"

    if not yaml_path.exists():
        # Если YAML нет — создаём пример
        print(f"⚠ Файл не найден: {yaml_path}")
        print("  Создаю пример из examples/sample_problem_list.yaml")
        example = Path(__file__).parent / "examples" / "sample_problem_list.yaml"
        if example.exists():
            yaml_path = example
        else:
            print("  Пример тоже не найден. Запустите сначала: python analyze.py plan")
            sys.exit(1)

    plan = load_plan(yaml_path)
    md_text = render_markdown(plan)

    # Сохраняем MD
    out_md = yaml_path.with_suffix(".md")
    out_md.write_text(md_text, encoding="utf-8")
    print(f"✓ Markdown: {out_md}")

    # Сохраняем DOCX
    out_docx = yaml_path.with_suffix(".docx")
    render_docx(md_text, out_docx)
